import os
import io
import re
from docx import Document
from google.cloud import storage
from google.api_core import exceptions

def create_docx_in_memory(document_title: str, document_content: str) -> io.BytesIO:
    """Helper function to create a .docx file in memory from markdown."""
    document = Document()
    document.add_heading(document_title, level=1)
    
    # A simple parser to convert markdown to docx formatting
    for line in document_content.split('\n'):
        if line.startswith('### '):
            document.add_heading(line.replace('### ', ''), level=3)
        elif line.startswith('## '):
            document.add_heading(line.replace('## ', ''), level=2)
        elif line.startswith('- **'):
            clean_line = re.sub(r'^\s*-\s*\*\*(.*?)\*\*:', r'\1:', line)
            parts = clean_line.split(':', 1)
            p = document.add_paragraph(style='List Bullet')
            p.add_run(parts[0].strip()).bold = True
            if len(parts) > 1:
                p.add_run(f":{parts[1].strip()}")
        elif line.strip() == "":
            continue
        else:
            document.add_paragraph(line)

    doc_stream = io.BytesIO()
    document.save(doc_stream)
    doc_stream.seek(0)
    return doc_stream

def upload_docx_to_gcs(document_title: str, document_content: str, bucket_name: str, destination_folder: str) -> str:
    """
    Creates a .docx file from text and uploads it to a Google Cloud Storage bucket.

    Args:
        document_title: The title for the document, used as the base for the filename.
        document_content: The full text content of the document, in markdown format.
        bucket_name: The name of the GCS bucket to upload to (e.g., 'adk-agent-test').
        destination_folder: The folder path within the bucket (e.g., 'DAT').

    Returns:
        A string containing the GCS URI of the created object or an error message.
    """
    if not bucket_name:
        return "Error: bucket_name parameter is missing."

    try:
        storage_client = storage.Client()
        doc_stream = create_docx_in_memory(document_title, document_content)
        
        bucket = storage_client.bucket(bucket_name)
        file_name = f"{document_title.replace(' ', '_')}.docx"
        destination_blob_name = f"{destination_folder}/{file_name}" if destination_folder else file_name
        blob = bucket.blob(destination_blob_name)

        blob.upload_from_file(doc_stream, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

        return f"Successfully exported document to Google Cloud Storage: gs://{bucket_name}/{destination_blob_name}"
    except exceptions.NotFound:
        return f"An error occurred: The bucket '{bucket_name}' does not exist."
    except exceptions.Forbidden as e:
        return f"An error occurred: Permission denied for bucket '{bucket_name}'. Ensure the account has 'Storage Object Creator' role. Details: {str(e)}"
    except Exception as e:
        return f"An unexpected error occurred while uploading to GCS: {str(e)}"